# file_utils/word.py
"""
Модуль для вставки штампа "Оплачено ДатаЧас" у Word-документи (.docx та .doc).
- .docx: вставка першого рядка з `python-docx`
- .doc: через pywin32 (MS Word має бути встановлений)
"""

import os
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from log import log

def stamp_doc(file_path: str, stamp: str):
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".docx":
            _stamp_docx(file_path, stamp)
        elif ext == ".doc":
            _stamp_doc_legacy(file_path, stamp)
        else:
            log(f"⚠️ Непідтримуваний формат Word: {file_path}")
    except Exception as e:
        log(f"❌ Помилка при вставці штампа в Word: {file_path} — {e}")


def _stamp_docx(file_path: str, stamp: str):
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(file_path)

    # Створюємо новий параграф на початку
    para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
    run = para.insert_paragraph_before().add_run("Paid " + stamp.replace("Оплачено ", ""))
    run.bold = True
    run.font.size = Pt(14)

    doc.save(file_path)
    log(f"✅ DOCX оброблено: {file_path}")


def _stamp_doc_legacy(file_path: str, stamp: str):
    import win32com.client as win32
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False

    doc = word.Documents.Open(file_path)
    selection = word.Selection
    selection.HomeKey(Unit=6)  # WdStory
    selection.TypeText(f"Paid {stamp.replace('Оплачено ', '')}\n")

    doc.Save()
    doc.Close()
    word.Quit()
    log(f"✅ DOC оброблено: {file_path}")


# 🔧 Тест запуску незалежно
if __name__ == "__main__":
    stamp_doc(r"C:\Users\la\OneDrive\Рабочий стол\test.docx", "Оплачено 2025-07-30 20:40")
    stamp_doc(r"C:\Users\la\OneDrive\Рабочий стол\test.doc", "Оплачено 2025-07-30 20:40")
